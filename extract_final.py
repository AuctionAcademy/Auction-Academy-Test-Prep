from docx import Document
import json

def extract_tn_questions_final(filename):
    """Extract questions from TN - each question/answer/explanation is sequential"""
    doc = Document(filename)
    questions = []
    current_question = None
    current_topic = ""
    state = "looking_for_question"  # states: looking_for_question, reading_options, reading_answer, reading_explanation
    
    i = 0
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        if not text:
            i += 1
            continue
        
        # Topic line
        if text.startswith("Topic:"):
            current_topic = text.replace("Topic:", "").strip()
            state = "looking_for_question"
            i += 1
            continue
        
        # Looking for a question (ends with ?)
        if state == "looking_for_question" and text.endswith("?"):
            # Save previous question if exists
            if current_question and len(current_question.get("options", [])) == 4:
                questions.append(current_question)
            
            # Start new question
            current_question = {
                "topic": current_topic or "General",
                "question": text,
                "options": [],
                "correctAnswer": -1,
                "explanation": ""
            }
            state = "reading_options"
            i += 1
            continue
        
        # Reading options (A., B., C., D.)
        if state == "reading_options":
            if text.startswith("A. "):
                current_question["options"].append(text[3:])
            elif text.startswith("B. "):
                current_question["options"].append(text[3:])
            elif text.startswith("C. "):
                current_question["options"].append(text[3:])
            elif text.startswith("D. "):
                current_question["options"].append(text[3:])
                # After D, next should be Answer
                state = "reading_answer"
            i += 1
            continue
        
        # Reading answer
        if state == "reading_answer" and text.startswith("Answer:"):
            answer_letter = text.replace("Answer:", "").strip().upper()
            if answer_letter.startswith("A"):
                current_question["correctAnswer"] = 0
            elif answer_letter.startswith("B"):
                current_question["correctAnswer"] = 1
            elif answer_letter.startswith("C"):
                current_question["correctAnswer"] = 2
            elif answer_letter.startswith("D"):
                current_question["correctAnswer"] = 3
            state = "reading_explanation"
            i += 1
            continue
        
        # Reading explanation
        if state == "reading_explanation":
            if text.startswith("Explanation:"):
                expl_text = text.replace("Explanation:", "").strip()
                if expl_text:
                    current_question["explanation"] = expl_text
            elif text.startswith("Topic:"):
                # New topic, go back
                current_topic = text.replace("Topic:", "").strip()
                state = "looking_for_question"
            elif text.endswith("?"):
                # New question without topic change
                state = "looking_for_question"
                continue  # Don't increment i, process this line again
            else:
                # Continue explanation
                if current_question["explanation"]:
                    current_question["explanation"] += " " + text
                else:
                    current_question["explanation"] = text
        
        i += 1
    
    # Save last question
    if current_question and len(current_question.get("options", [])) == 4:
        questions.append(current_question)
    
    return questions

# Extract TN questions
print("Extracting TN questions (final version)...")
tn_questions = extract_tn_questions_final("Tennessee Auctioneer Exam Bank (2).docx")
print(f"Extracted {len(tn_questions)} TN questions")

with open("tn_questions_updated.json", "w") as f:
    json.dump(tn_questions, f, indent=2)

# Show topic distribution
tn_topics = {}
for q in tn_questions:
    topic = q.get("topic", "Unknown")
    tn_topics[topic] = tn_topics.get(topic, 0) + 1

print("\n=== TN Topic Distribution ===")
for topic, count in sorted(tn_topics.items(), key=lambda x: x[1], reverse=True):
    print(f"{topic}: {count} questions")

# Show samples
print("\n=== Sample TN Questions ===")
for i in [0, 10, 50, 100, len(tn_questions)-1]:
    if i < len(tn_questions):
        q = tn_questions[i]
        print(f"\nQuestion {i+1}:")
        print(f"Topic: {q['topic']}")
        print(f"Q: {q['question'][:70]}...")
        print(f"Options: {len(q['options'])}")
        print(f"Answer: {chr(65 + q['correctAnswer']) if q['correctAnswer'] >= 0 else 'N/A'}")
        print(f"Expl: {q['explanation'][:60]}...")

print(f"\n✅ Total: {len(tn_questions)} TN questions extracted")

