#!/usr/bin/env python3
"""
Extract Texas questions with correct answers, topics, and clean text.
Fixes issues:
- Removes "Question #." prefix
- Correctly parses answers (B = index 1, not 0)
- Assigns appropriate topics based on content
"""

from docx import Document
import json
import re

def categorize_question(question_text):
    """Determine topic based on question content"""
    q_lower = question_text.lower()
    
    # Auction Math patterns
    if any(word in q_lower for word in ['hammer price', 'buyer premium', 'commission', 'calculate', 
                                          'sales tax', 'total due', 'percentage', 'profit', 'price']):
        return "Auction Math"
    
    # Licensing patterns
    if any(word in q_lower for word in ['license', 'licensing', 'apprentice', 'registration fee']):
        return "Licensing Requirements"
    
    # Ethics patterns
    if any(word in q_lower for word in ['ethical', 'ethics', 'conflict of interest', 'fiduciary', 
                                          'duty', 'professional conduct', 'integrity']):
        return "Ethics and Professional Conduct"
    
    # Record keeping patterns
    if any(word in q_lower for word in ['record', 'documentation', 'settlement', 'keep records',
                                          'accountability', 'audit', 'invoice']):
        return "Record Keeping"
    
    # State-specific laws
    if any(word in q_lower for word in ['texas', 'statute', 'regulation', 'occupations code',
                                          'business and commerce code', 'state law']):
        return "State-Specific Laws"
    
    # Bidding procedures
    if any(word in q_lower for word in ['bidder', 'bidding', 'bid', 'reserve', 'increment',
                                          'absolute auction', 'with reserve', 'without reserve']):
        return "Bidding Procedures"
    
    # Contract law
    if any(word in q_lower for word in ['contract', 'agreement', 'terms and conditions',
                                          'breach', 'enforceable']):
        return "Contract Law"
    
    # UCC
    if any(word in q_lower for word in ['ucc', 'uniform commercial code', 'secured transaction',
                                          'article 9', 'security interest']):
        return "UCC (Uniform Commercial Code)"
    
    # Advertising
    if any(word in q_lower for word in ['advertis', 'marketing', 'promotion', 'announce',
                                          'publication', 'disclosure']):
        return "Advertising and Marketing"
    
    # Consumer protection
    if any(word in q_lower for word in ['consumer protection', 'buyer protection', 'fraud',
                                          'misrepresentation', 'deceptive']):
        return "Consumer Protection"
    
    # Real estate
    if any(word in q_lower for word in ['real estate', 'property', 'land', 'foreclosure',
                                          'deed', 'title']):
        return "Real Estate Auctions"
    
    # Default to Auction Basics
    return "Auction Basics"

def extract_texas_questions(docx_path):
    """Extract all questions from Texas document with proper parsing"""
    doc = Document(docx_path)
    questions = []
    
    i = 0
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        # Look for "Question N." pattern
        if text.startswith('Question ') and '. ' in text:
            # Extract question number and text
            match = re.match(r'Question (\d+)\.\s*(.*)', text)
            if match:
                q_num = int(match.group(1))
                question_text = match.group(2).strip()
                
                # Get options (A, B, C, D)
                options = []
                correct_answer = None
                j = i + 1
                
                # Look for options and answer
                while j < len(doc.paragraphs) and j < i + 10:
                    para_text = doc.paragraphs[j].text.strip()
                    
                    if para_text.startswith('A. '):
                        options.append(para_text[3:].strip())
                    elif para_text.startswith('B. '):
                        options.append(para_text[3:].strip())
                    elif para_text.startswith('C. '):
                        options.append(para_text[3:].strip())
                    elif para_text.startswith('D. '):
                        options.append(para_text[3:].strip())
                    elif para_text.startswith('Answer:'):
                        answer_letter = para_text.split(':')[1].strip().upper()
                        if answer_letter == 'A':
                            correct_answer = 0
                        elif answer_letter == 'B':
                            correct_answer = 1
                        elif answer_letter == 'C':
                            correct_answer = 2
                        elif answer_letter == 'D':
                            correct_answer = 3
                        break
                    
                    j += 1
                
                # Only add if we have all components
                if len(options) == 4 and correct_answer is not None:
                    topic = categorize_question(question_text)
                    
                    # Create explanation
                    answer_letter = ['A', 'B', 'C', 'D'][correct_answer]
                    explanation = f"The correct answer is {answer_letter}. {options[correct_answer]}"
                    
                    questions.append({
                        "id": q_num,
                        "topic": topic,
                        "question": question_text,  # No "Question #." prefix
                        "options": options,
                        "correctAnswer": correct_answer,
                        "explanation": explanation
                    })
                    
                    if q_num % 50 == 0:
                        print(f"Processed {q_num} questions...")
        
        i += 1
    
    return questions

# Extract questions
print("Extracting Texas questions from source document...")
questions = extract_texas_questions('tx Practice Questions and study guide (2).docx')

print(f"\nExtracted {len(questions)} questions")

# Show topic distribution
topics = {}
for q in questions:
    topic = q['topic']
    topics[topic] = topics.get(topic, 0) + 1

print("\nTopic distribution:")
for topic, count in sorted(topics.items(), key=lambda x: x[1], reverse=True):
    print(f"  {topic}: {count}")

# Save to JSON
with open('texas_questions_fixed.json', 'w') as f:
    json.dump(questions, f, indent=2)

print("\nSaved to texas_questions_fixed.json")

# Show first few questions
print("\nFirst 3 questions:")
for q in questions[:3]:
    print(f"\nQ{q['id']}: [{q['topic']}]")
    print(f"  Question: {q['question'][:80]}...")
    print(f"  Options: {q['options']}")
    print(f"  Correct: {q['correctAnswer']} ({['A','B','C','D'][q['correctAnswer']]})")
    print(f"  Explanation: {q['explanation'][:80]}...")
