from docx import Document
import json
import re

def extract_tn_all_questions(filename):
    """Extract ALL questions from TN question bank"""
    doc = Document(filename)
    questions = []
    current_question = None
    current_topic = ""
    question_buffer = []
    explanation_buffer = []
    in_options = False
    in_explanation = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Topic line
        if text.startswith("Topic:"):
            # Save previous question before changing topic
            if current_question and current_question.get("correctAnswer") >= 0:
                if explanation_buffer:
                    current_question["explanation"] = " ".join(explanation_buffer).strip()
                if len(current_question.get("options", [])) == 4:
                    questions.append(current_question)
            
            current_topic = text.replace("Topic:", "").strip()
            current_question = None
            question_buffer = []
            explanation_buffer = []
            in_options = False
            in_explanation = False
            continue
        
        # Answer line - marks end of options and start of answer/explanation
        if text.startswith("Answer:"):
            # Finalize question text
            if question_buffer and current_question:
                current_question["question"] = " ".join(question_buffer).strip()
                question_buffer = []
            
            # Get answer
            answer_letter = text.replace("Answer:", "").strip().upper()
            if current_question:
                if answer_letter.startswith("A"):
                    current_question["correctAnswer"] = 0
                elif answer_letter.startswith("B"):
                    current_question["correctAnswer"] = 1
                elif answer_letter.startswith("C"):
                    current_question["correctAnswer"] = 2
                elif answer_letter.startswith("D"):
                    current_question["correctAnswer"] = 3
            
            in_options = False
            in_explanation = False
            continue
        
        # Explanation line
        if text.startswith("Explanation:"):
            in_explanation = True
            expl_text = text.replace("Explanation:", "").strip()
            if expl_text:
                explanation_buffer.append(expl_text)
            continue
        
        # Option lines
        if text.startswith("A. "):
            # This starts options, so finalize question text
            if question_buffer:
                if not current_question:
                    current_question = {
                        "topic": current_topic or "General",
                        "question": "",
                        "options": [],
                        "correctAnswer": -1,
                        "explanation": ""
                    }
                current_question["question"] = " ".join(question_buffer).strip()
                question_buffer = []
            
            if current_question:
                current_question["options"] = [text[3:]]
                in_options = True
            continue
            
        if text.startswith("B. ") and current_question and in_options:
            current_question["options"].append(text[3:])
            continue
            
        if text.startswith("C. ") and current_question and in_options:
            current_question["options"].append(text[3:])
            continue
            
        if text.startswith("D. ") and current_question and in_options:
            current_question["options"].append(text[3:])
            continue
        
        # Regular text
        if in_explanation:
            explanation_buffer.append(text)
        elif not in_options:
            # This is part of the question
            question_buffer.append(text)
    
    # Save last question
    if current_question and current_question.get("correctAnswer") >= 0:
        if explanation_buffer:
            current_question["explanation"] = " ".join(explanation_buffer).strip()
        if len(current_question.get("options", [])) == 4:
            questions.append(current_question)
    
    return questions

def extract_tx_all_questions(filename):
    """Extract ALL questions from TX question bank"""
    doc = Document(filename)
    questions = []
    current_question = None
    current_volume = ""
    question_buffer = []
    explanation_buffer = []
    in_options = False
    in_explanation = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Volume/Topic line
        if text.startswith("Volume") and ":" in text:
            current_volume = text.split("(")[0].strip() if "(" in text else text
            continue
        
        # Question line (starts with "Question #.")
        if re.match(r'^Question \d+\.', text):
            # Save previous question
            if current_question and current_question.get("correctAnswer") >= 0:
                if question_buffer:
                    current_question["question"] = " ".join(question_buffer).strip()
                if explanation_buffer:
                    current_question["explanation"] = " ".join(explanation_buffer).strip()
                if len(current_question.get("options", [])) == 4:
                    questions.append(current_question)
            
            # Extract question text (remove "Question #." prefix)
            question_text = re.sub(r'^Question \d+\.\s*', '', text)
            
            # Determine topic from volume
            topic = "Auction Basics"
            if current_volume:
                if "Auction Math" in current_volume or "Math" in current_volume:
                    topic = "Auction Math"
                elif "Texas Law" in current_volume or "State" in current_volume or "Law" in current_volume:
                    topic = "State-Specific Laws"
                elif "UCC" in current_volume or "Secured" in current_volume:
                    topic = "UCC"
                elif "Ethics" in current_volume:
                    topic = "Ethics and Professional Conduct"
            
            current_question = {
                "topic": topic,
                "question": question_text,
                "options": [],
                "correctAnswer": -1,
                "explanation": ""
            }
            question_buffer = []
            explanation_buffer = []
            in_options = False
            in_explanation = False
            continue
        
        # Answer line
        if text.startswith("Answer:"):
            answer_letter = text.replace("Answer:", "").strip().upper()
            if current_question:
                if answer_letter.startswith("A"):
                    current_question["correctAnswer"] = 0
                elif answer_letter.startswith("B"):
                    current_question["correctAnswer"] = 1
                elif answer_letter.startswith("C"):
                    current_question["correctAnswer"] = 2
                elif answer_letter.startswith("D"):
                    current_question["correctAnswer"] = 3
            in_options = False
            in_explanation = False
            continue
        
        # Explanation line
        if text.startswith("Explanation:"):
            in_explanation = True
            expl_text = text.replace("Explanation:", "").strip()
            if expl_text:
                explanation_buffer.append(expl_text)
            continue
        
        # Option lines
        if text.startswith("A. "):
            if current_question:
                current_question["options"] = [text[3:]]
                in_options = True
            continue
            
        if text.startswith("B. ") and current_question and in_options:
            current_question["options"].append(text[3:])
            continue
            
        if text.startswith("C. ") and current_question and in_options:
            current_question["options"].append(text[3:])
            continue
            
        if text.startswith("D. ") and current_question and in_options:
            current_question["options"].append(text[3:])
            continue
        
        # Regular text
        if in_explanation:
            explanation_buffer.append(text)
        elif current_question and not in_options:
            question_buffer.append(text)
    
    # Save last question
    if current_question and current_question.get("correctAnswer") >= 0:
        if question_buffer:
            current_question["question"] = " ".join(question_buffer).strip()
        if explanation_buffer:
            current_question["explanation"] = " ".join(explanation_buffer).strip()
        if len(current_question.get("options", [])) == 4:
            questions.append(current_question)
    
    return questions

# Extract TN questions
print("Extracting TN questions...")
tn_questions = extract_tn_all_questions("Tennessee Auctioneer Exam Bank (2).docx")
print(f"Extracted {len(tn_questions)} TN questions")

with open("tn_questions_updated.json", "w") as f:
    json.dump(tn_questions, f, indent=2)
print("Saved to tn_questions_updated.json")

# Extract TX questions
print("\nExtracting TX questions...")
tx_questions = extract_tx_all_questions("tx Practice Questions and study guide (5).docx")
print(f"Extracted {len(tx_questions)} TX questions")

with open("tx_questions_updated.json", "w") as f:
    json.dump(tx_questions, f, indent=2)
print("Saved to tx_questions_updated.json")

# Show stats
print(f"\n=== Total Questions ===")
print(f"TN: {len(tn_questions)} questions")
print(f"TX: {len(tx_questions)} questions")
print(f"Total: {len(tn_questions) + len(tx_questions)} questions")

# Show TN topic distribution
tn_topics = {}
for q in tn_questions:
    topic = q.get("topic", "Unknown")
    tn_topics[topic] = tn_topics.get(topic, 0) + 1

print("\n=== TN Topic Distribution ===")
for topic, count in sorted(tn_topics.items(), key=lambda x: x[1], reverse=True):
    print(f"{topic}: {count} questions")

# Show TX topic distribution
tx_topics = {}
for q in tx_questions:
    topic = q.get("topic", "Unknown")
    tx_topics[topic] = tx_topics.get(topic, 0) + 1

print("\n=== TX Topic Distribution ===")
for topic, count in sorted(tx_topics.items(), key=lambda x: x[1], reverse=True):
    print(f"{topic}: {count} questions")

