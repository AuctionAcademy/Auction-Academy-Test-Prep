from docx import Document
import json

def extract_tn_complete(filename):
    """Extract ALL TN questions - complete version"""
    doc = Document(filename)
    questions = []
    
    i = 0
    current_topic = ""
    
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        if not text:
            i += 1
            continue
        
        # Topic line
        if text.startswith("Topic:"):
            current_topic = text.replace("Topic:", "").strip()
            i += 1
            continue
        
        # Check if this looks like a question (ends with ?)
        if "?" in text:
            question_text = text
            options = []
            answer_idx = -1
            explanation = ""
            
            # Look ahead for options (A, B, C, D)
            j = i + 1
            while j < len(doc.paragraphs) and j < i + 10:  # Look ahead max 10 lines
                next_text = doc.paragraphs[j].text.strip()
                
                if next_text.startswith("A. "):
                    options.append(next_text[3:])
                    j += 1
                elif next_text.startswith("B. "):
                    options.append(next_text[3:])
                    j += 1
                elif next_text.startswith("C. "):
                    options.append(next_text[3:])
                    j += 1
                elif next_text.startswith("D. "):
                    options.append(next_text[3:])
                    j += 1
                    break  # Got all 4 options
                else:
                    j += 1
            
            # If we got 4 options, look for Answer
            if len(options) == 4:
                j += 1
                if j < len(doc.paragraphs):
                    answer_text = doc.paragraphs[j].text.strip()
                    if answer_text.startswith("Answer:"):
                        answer_letter = answer_text.replace("Answer:", "").strip().upper()
                        if answer_letter.startswith("A"):
                            answer_idx = 0
                        elif answer_letter.startswith("B"):
                            answer_idx = 1
                        elif answer_letter.startswith("C"):
                            answer_idx = 2
                        elif answer_letter.startswith("D"):
                            answer_idx = 3
                        
                        # Look for Explanation
                        j += 1
                        expl_lines = []
                        while j < len(doc.paragraphs) and j < i + 20:
                            expl_text = doc.paragraphs[j].text.strip()
                            if expl_text.startswith("Explanation:"):
                                expl_text = expl_text.replace("Explanation:", "").strip()
                                if expl_text:
                                    expl_lines.append(expl_text)
                                j += 1
                            elif expl_text and not expl_text.startswith("Topic:") and not "?" in expl_text:
                                expl_lines.append(expl_text)
                                j += 1
                            else:
                                break
                        
                        explanation = " ".join(expl_lines)
                        
                        # Save question
                        if answer_idx >= 0:
                            questions.append({
                                "topic": current_topic or "General",
                                "question": question_text,
                                "options": options,
                                "correctAnswer": answer_idx,
                                "explanation": explanation
                            })
                        
                        # Skip to after this question
                        i = j
                        continue
        
        i += 1
    
    return questions

def extract_tx_complete(filename):
    """Extract ALL TX questions"""
    doc = Document(filename)
    questions = []
    
    i = 0
    current_topic = "Auction Basics"
    
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        if not text:
            i += 1
            continue
        
        # Volume line determines topic
        if text.startswith("Volume"):
            if "Auction Math" in text or "Math" in text:
                current_topic = "Auction Math"
            elif "Texas Law" in text or "State" in text or "Law" in text:
                current_topic = "State-Specific Laws"
            elif "UCC" in text or "Secured" in text:
                current_topic = "UCC"
            elif "Ethics" in text:
                current_topic = "Ethics and Professional Conduct"
            else:
                current_topic = "Auction Basics"
            i += 1
            continue
        
        # Question line (starts with "Question #.")
        import re
        if re.match(r'^Question \d+\.', text):
            # Remove Question prefix
            question_text = re.sub(r'^Question \d+\.\s*', '', text)
            options = []
            answer_idx = -1
            explanation = ""
            
            # Look ahead for options
            j = i + 1
            while j < len(doc.paragraphs) and j < i + 10:
                next_text = doc.paragraphs[j].text.strip()
                
                if next_text.startswith("A. "):
                    options.append(next_text[3:])
                    j += 1
                elif next_text.startswith("B. "):
                    options.append(next_text[3:])
                    j += 1
                elif next_text.startswith("C. "):
                    options.append(next_text[3:])
                    j += 1
                elif next_text.startswith("D. "):
                    options.append(next_text[3:])
                    j += 1
                    break
                else:
                    j += 1
            
            # Look for Answer
            if len(options) == 4:
                j += 1
                if j < len(doc.paragraphs):
                    answer_text = doc.paragraphs[j].text.strip()
                    if answer_text.startswith("Answer:"):
                        answer_letter = answer_text.replace("Answer:", "").strip().upper()
                        if answer_letter.startswith("A"):
                            answer_idx = 0
                        elif answer_letter.startswith("B"):
                            answer_idx = 1
                        elif answer_letter.startswith("C"):
                            answer_idx = 2
                        elif answer_letter.startswith("D"):
                            answer_idx = 3
                        
                        # Look for Explanation
                        j += 1
                        expl_lines = []
                        while j < len(doc.paragraphs) and j < i + 20:
                            expl_text = doc.paragraphs[j].text.strip()
                            if expl_text.startswith("Explanation:"):
                                expl_text = expl_text.replace("Explanation:", "").strip()
                                if expl_text:
                                    expl_lines.append(expl_text)
                                j += 1
                            elif expl_text and not expl_text.startswith(("Question", "Volume")):
                                expl_lines.append(expl_text)
                                j += 1
                            else:
                                break
                        
                        explanation = " ".join(expl_lines)
                        
                        # Save question
                        if answer_idx >= 0:
                            questions.append({
                                "topic": current_topic,
                                "question": question_text,
                                "options": options,
                                "correctAnswer": answer_idx,
                                "explanation": explanation
                            })
                        
                        i = j
                        continue
        
        i += 1
    
    return questions

# Extract both
print("Extracting TN questions (complete)...")
tn_questions = extract_tn_complete("Tennessee Auctioneer Exam Bank (2).docx")
print(f"✅ Extracted {len(tn_questions)} TN questions")

with open("tn_questions_updated.json", "w") as f:
    json.dump(tn_questions, f, indent=2)

print("\nExtracting TX questions (complete)...")
tx_questions = extract_tx_complete("tx Practice Questions and study guide (5).docx")
print(f"✅ Extracted {len(tx_questions)} TX questions")

with open("tx_questions_updated.json", "w") as f:
    json.dump(tx_questions, f, indent=2)

print(f"\n=== TOTAL ===")
print(f"TN: {len(tn_questions)} questions")
print(f"TX: {len(tx_questions)} questions")
print(f"Total: {len(tn_questions) + len(tx_questions)} questions")

# Topic distribution
tn_topics = {}
for q in tn_questions:
    t = q["topic"]
    tn_topics[t] = tn_topics.get(t, 0) + 1

print("\n=== TN Topics ===")
for topic, count in sorted(tn_topics.items(), key=lambda x: x[1], reverse=True):
    print(f"{topic}: {count}")

tx_topics = {}
for q in tx_questions:
    t = q["topic"]
    tx_topics[t] = tx_topics.get(t, 0) + 1

print("\n=== TX Topics ===")
for topic, count in sorted(tx_topics.items(), key=lambda x: x[1], reverse=True):
    print(f"{topic}: {count}")

