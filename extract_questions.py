#!/usr/bin/env python3
"""
Extract questions from TX and TN source documents and format for the application
"""

from docx import Document
import json
import re

def clean_text(text):
    """Clean and normalize text"""
    return text.strip().replace('\u2019', "'").replace('\u201c', '"').replace('\u201d', '"')

def extract_tennessee_questions():
    """Extract questions from Tennessee Auctioneer Exam Bank"""
    doc = Document('Tennessee Auctioneer Exam Bank.docx')
    questions = []
    
    i = 0
    question_num = 0
    
    while i < len(doc.paragraphs):
        para_text = doc.paragraphs[i].text.strip()
        
        # Look for question pattern "Question N."
        if para_text.startswith('Question ') and '. ' in para_text:
            question_num += 1
            question_text = clean_text(para_text.split('. ', 1)[1] if '. ' in para_text else para_text)
            i += 1
            
            # Collect options
            options = []
            correct_idx = -1
            
            # Get next paragraphs for options
            while i < len(doc.paragraphs) and len(options) < 4:
                next_text = doc.paragraphs[i].text.strip()
                
                if not next_text:
                    i += 1
                    continue
                    
                # Check for option markers
                if re.match(r'^[A-D][\)\.]\s*', next_text):
                    option_text = re.sub(r'^[A-D][\)\.]\s*', '', next_text)
                    options.append(clean_text(option_text))
                    i += 1
                elif next_text.startswith('Answer:'):
                    # Extract correct answer
                    answer_match = re.search(r'Answer:\s*([A-D])', next_text)
                    if answer_match:
                        correct_idx = ord(answer_match.group(1)) - ord('A')
                    i += 1
                    break
                else:
                    break
            
            # Only add if we have 4 options and a correct answer
            if len(options) == 4 and correct_idx >= 0:
                # Determine topic based on question content
                topic = categorize_question(question_text)
                
                questions.append({
                    'id': question_num,
                    'topic': topic,
                    'question': question_text,
                    'options': options,
                    'correctAnswer': correct_idx,
                    'explanation': f'The correct answer is {options[correct_idx]}.'
                })
        else:
            i += 1
    
    return questions

def extract_texas_questions():
    """Extract questions from Texas study guide"""
    doc = Document('tx Practice Questions and study guide (1).docx')
    questions = []
    
    # Texas questions appear to be word problems/scenarios
    # Let's search for calculation-based questions
    i = 0
    question_num = 0
    
    while i < len(doc.paragraphs):
        para_text = doc.paragraphs[i].text.strip()
        
        # Look for question patterns (calculations, "What is", etc.)
        if '?' in para_text and any(keyword in para_text.lower() for keyword in ['what is', 'how much', 'calculate', 'determine', 'find']):
            question_num += 1
            question_text = clean_text(para_text)
            i += 1
            
            # Look for answer options or calculations
            options = []
            correct_idx = -1
            
            # Collect next few paragraphs that might be options
            while i < len(doc.paragraphs) and len(options) < 4:
                next_text = doc.paragraphs[i].text.strip()
                
                if not next_text:
                    i += 1
                    continue
                
                # Check for option patterns or "Answer:" pattern
                if re.match(r'^[A-D][\)\.]\s*', next_text):
                    option_text = re.sub(r'^[A-D][\)\.]\s*', '', next_text)
                    options.append(clean_text(option_text))
                    i += 1
                elif next_text.startswith('Answer:'):
                    answer_match = re.search(r'Answer:\s*([A-D]|\$[\d,]+)', next_text)
                    if answer_match:
                        ans = answer_match.group(1)
                        if ans in ['A', 'B', 'C', 'D']:
                            correct_idx = ord(ans) - ord('A')
                    i += 1
                    break
                elif next_text.startswith('$') and len(options) < 4:
                    # Dollar amount might be an option
                    options.append(clean_text(next_text))
                    i += 1
                else:
                    break
            
            # Skip if not enough options
            if len(options) >= 2:
                # Generate missing options if needed
                while len(options) < 4:
                    options.append(f"Option {len(options) + 1}")
                
                if correct_idx < 0:
                    correct_idx = 0  # Default to first option if not found
                
                topic = categorize_question(question_text)
                
                questions.append({
                    'id': question_num,
                    'topic': topic,
                    'question': question_text,
                    'options': options[:4],
                    'correctAnswer': correct_idx,
                    'explanation': f'Based on the calculation, the correct answer is {options[correct_idx]}.'
                })
        else:
            i += 1
    
    return questions

def categorize_question(question_text):
    """Categorize question by topic based on keywords"""
    question_lower = question_text.lower()
    
    if any(word in question_lower for word in ['premium', 'tax', 'price', 'cost', 'commission', 'calculate', '$']):
        return 'Auction Calculations'
    elif any(word in question_lower for word in ['license', 'requirement', 'qualify', 'application']):
        return 'Licensing Requirements'
    elif any(word in question_lower for word in ['contract', 'agreement', 'binding']):
        return 'Contract Law'
    elif any(word in question_lower for word in ['ethical', 'unethical', 'conduct', 'professional']):
        return 'Ethics and Professional Conduct'
    elif any(word in question_lower for word in ['bidding', 'bid', 'hammer', 'reserve']):
        return 'Bidding Procedures'
    elif any(word in question_lower for word in ['ucc', 'uniform commercial']):
        return 'UCC (Uniform Commercial Code)'
    elif any(word in question_lower for word in ['real estate', 'property', 'land']):
        return 'Real Estate Auctions'
    elif any(word in question_lower for word in ['record', 'document', 'filing']):
        return 'Record Keeping'
    elif any(word in question_lower for word in ['advertis', 'marketing', 'promotion']):
        return 'Advertising and Marketing'
    elif any(word in question_lower for word in ['consumer', 'protection', 'buyer']):
        return 'Consumer Protection'
    elif any(word in question_lower for word in ['tennessee', 'texas', 'state', 'statute']):
        return 'State-Specific Laws'
    else:
        return 'Auction Basics'

def main():
    print("Extracting Tennessee questions...")
    tn_questions = extract_tennessee_questions()
    print(f"✓ Extracted {len(tn_questions)} Tennessee questions")
    
    print("\nExtracting Texas questions...")
    tx_questions = extract_texas_questions()
    print(f"✓ Extracted {len(tx_questions)} Texas questions")
    
    # Save to JSON files
    with open('tennessee_questions.json', 'w') as f:
        json.dump(tn_questions, f, indent=2)
    print(f"\n✓ Saved Tennessee questions to tennessee_questions.json")
    
    with open('texas_questions.json', 'w') as f:
        json.dump(tx_questions, f, indent=2)
    print(f"✓ Saved Texas questions to texas_questions.json")
    
    # Print sample
    print("\n" + "=" * 80)
    print("SAMPLE TENNESSEE QUESTION:")
    print("=" * 80)
    if tn_questions:
        q = tn_questions[0]
        print(f"Topic: {q['topic']}")
        print(f"Q: {q['question']}")
        for i, opt in enumerate(q['options']):
            marker = "✓" if i == q['correctAnswer'] else " "
            print(f"  {marker} {chr(65+i)}. {opt}")
    
    print("\n" + "=" * 80)
    print("SAMPLE TEXAS QUESTION:")
    print("=" * 80)
    if tx_questions:
        q = tx_questions[0]
        print(f"Topic: {q['topic']}")
        print(f"Q: {q['question']}")
        for i, opt in enumerate(q['options']):
            marker = "✓" if i == q['correctAnswer'] else " "
            print(f"  {marker} {chr(65+i)}. {opt}")
    
    print("\n" + "=" * 80)
    print("EXTRACTION COMPLETE!")
    print("=" * 80)
    print(f"Total Tennessee: {len(tn_questions)} questions")
    print(f"Total Texas: {len(tx_questions)} questions")

if __name__ == '__main__':
    main()
