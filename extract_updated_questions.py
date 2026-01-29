from docx import Document
import json
import re

def extract_tn_questions(filename):
    """Extract questions from UPDATED TN question bank"""
    doc = Document(filename)
    questions = []
    current_question = {}
    current_topic = ""
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Topic line
        if text.startswith("Topic:"):
            current_topic = text.replace("Topic:", "").strip()
            continue
        
        # Check if this looks like a question (ends with ?)
        if text.endswith("?") and not text.startswith("A.") and not text.startswith("B.") and not text.startswith("C.") and not text.startswith("D."):
            # Save previous question if exists
            if current_question.get("question"):
                questions.append(current_question)
            
            current_question = {
                "topic": current_topic or "General",
                "question": text,
                "options": [],
                "correctAnswer": -1,
                "explanation": ""
            }
            continue
        
        # Option lines
        if text.startswith("A. "):
            current_question["options"].append(text[3:])
        elif text.startswith("B. "):
            current_question["options"].append(text[3:])
        elif text.startswith("C. "):
            current_question["options"].append(text[3:])
        elif text.startswith("D. "):
            current_question["options"].append(text[3:])
        
        # Answer line
        elif text.startswith("Answer:"):
            answer_letter = text.replace("Answer:", "").strip()
            if answer_letter == "A":
                current_question["correctAnswer"] = 0
            elif answer_letter == "B":
                current_question["correctAnswer"] = 1
            elif answer_letter == "C":
                current_question["correctAnswer"] = 2
            elif answer_letter == "D":
                current_question["correctAnswer"] = 3
        
        # Explanation line
        elif text.startswith("Explanation:"):
            current_question["explanation"] = text.replace("Explanation:", "").strip()
        elif current_question.get("explanation") and text and not text.startswith("Topic:"):
            # Continue explanation on next lines
            if len(current_question["explanation"]) > 0:
                current_question["explanation"] += " " + text
            else:
                current_question["explanation"] = text
    
    # Add last question
    if current_question.get("question"):
        questions.append(current_question)
    
    return questions

def extract_tx_questions(filename):
    """Extract questions from UPDATED TX question bank (file 5)"""
    doc = Document(filename)
    questions = []
    current_question = {}
    current_volume = ""
    in_explanation = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Volume/Topic line
        if text.startswith("Volume"):
            current_volume = text.split("(")[0].strip() if "(" in text else text
            continue
        
        # Question line (starts with "Question #.")
        if re.match(r'^Question \d+\.', text):
            # Save previous question if exists
            if current_question.get("question"):
                questions.append(current_question)
            
            # Extract question text (remove "Question #." prefix)
            question_text = re.sub(r'^Question \d+\.\s*', '', text)
            
            # Determine topic from volume
            topic = "Auction Basics"
            if "Auction Math" in current_volume or "Math" in current_volume:
                topic = "Auction Math"
            elif "Texas Law" in current_volume or "State" in current_volume:
                topic = "State-Specific Laws"
            elif "UCC" in current_volume or "Secured" in current_volume:
                topic = "UCC"
            elif "Ethics" in current_volume:
                topic = "Ethics and Professional Conduct"
            
            current_question = {
                "topic": topic,
                "question": question_text,
                "options": [],
                "correctAnswer": -1,
                "explanation": ""
            }
            in_explanation = False
            continue
        
        # Option lines
        if text.startswith("A. "):
            current_question["options"].append(text[3:])
        elif text.startswith("B. "):
            current_question["options"].append(text[3:])
        elif text.startswith("C. "):
            current_question["options"].append(text[3:])
        elif text.startswith("D. "):
            current_question["options"].append(text[3:])
        
        # Answer line
        elif text.startswith("Answer:"):
            answer_letter = text.replace("Answer:", "").strip()
            if answer_letter == "A":
                current_question["correctAnswer"] = 0
            elif answer_letter == "B":
                current_question["correctAnswer"] = 1
            elif answer_letter == "C":
                current_question["correctAnswer"] = 2
            elif answer_letter == "D":
                current_question["correctAnswer"] = 3
        
        # Explanation line
        elif text.startswith("Explanation:"):
            in_explanation = True
            current_question["explanation"] = text.replace("Explanation:", "").strip()
        elif in_explanation and text and not text.startswith("Question"):
            # Continue explanation
            if current_question["explanation"]:
                current_question["explanation"] += " " + text
            else:
                current_question["explanation"] = text
    
    # Add last question
    if current_question.get("question"):
        questions.append(current_question)
    
    return questions

# Extract TN questions
print("Extracting TN questions...")
tn_questions = extract_tn_questions("Tennessee Auctioneer Exam Bank (2).docx")
print(f"Extracted {len(tn_questions)} TN questions")

# Save to JSON
with open("tn_questions_updated.json", "w") as f:
    json.dump(tn_questions, f, indent=2)
print("Saved to tn_questions_updated.json")

# Extract TX questions
print("\nExtracting TX questions...")
tx_questions = extract_tx_questions("tx Practice Questions and study guide (5).docx")
print(f"Extracted {len(tx_questions)} TX questions")

# Save to JSON
with open("tx_questions_updated.json", "w") as f:
    json.dump(tx_questions, f, indent=2)
print("Saved to tx_questions_updated.json")

# Show samples
print("\n=== Sample TN Question ===")
if tn_questions:
    q = tn_questions[0]
    print(f"Topic: {q['topic']}")
    print(f"Question: {q['question'][:100]}...")
    print(f"Options: {len(q['options'])}")
    print(f"Answer: {q['correctAnswer']}")
    print(f"Explanation: {q['explanation'][:100]}...")

print("\n=== Sample TX Question ===")
if tx_questions:
    q = tx_questions[0]
    print(f"Topic: {q['topic']}")
    print(f"Question: {q['question'][:100]}...")
    print(f"Options: {len(q['options'])}")
    print(f"Answer: {q['correctAnswer']}")
    print(f"Explanation: {q['explanation'][:100]}...")

